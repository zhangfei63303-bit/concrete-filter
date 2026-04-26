#!/usr/bin/env python3
"""
混凝土试件记录筛选工具 - GUI版本
支持：单文件/文件夹批量处理，递归子文件夹扫描
龄期筛选（支持直接读龄期列，或从制作日期计算）
输出：Word文档格式筛选结果
"""

import os
import sys
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime
import pandas as pd
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ========== 默认筛选参数 ==========
DEFAULT_AGE_MIN = 40
DEFAULT_AGE_MAX = 50
DEFAULT_DATE_START = ""
DEFAULT_DATE_END = ""
DEFAULT_SITE_NAME = ""
DEFAULT_GRADE = ""
DEFAULT_POUR_PART = ""


def parse_date(s):
    if not s or str(s).strip() == "":
        return None
    s = str(s).strip()
    for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%Y%m%d", "%m/%d/%Y", "%m.%d", "%m/%d"]:
        try:
            dt = datetime.strptime(s, fmt)
            if dt.year == 1900:
                dt = dt.replace(year=datetime.now().year)
            return dt
        except:
            continue
    return None


def calc_age(make_str):
    make_date = parse_date(make_str)
    if not make_date:
        return None
    today = datetime.now()
    age = (today - make_date).days
    return age if age > 0 else None


def read_excel(path):
    return pd.read_excel(path)


def read_csv(path):
    return pd.read_csv(path)


def read_docx(path):
    """读取 .docx 文件，智能识别表头行"""
    doc = Document(path)
    all_tables_data = []
    for table in doc.tables:
        for row in table.rows:
            all_tables_data.append([cell.text.strip() for cell in row.cells])

    if not all_tables_data:
        return None

    # 智能找表头行：在所有行中搜索包含"工程部位"或"龄期"的行
    header_row_idx = None
    for idx, row in enumerate(all_tables_data):
        first_cell = row[0].strip()
        if not first_cell or first_cell in ["", "工程名称", "检测日期"]:
            continue
        if any("工 程 部 位" in cell or "龄期" in cell for cell in row[:8]):
            header_row_idx = idx
            break

    if header_row_idx is None:
        return None

    header = [c.replace(" ", "").replace("\n", "").strip() for c in all_tables_data[header_row_idx]]

    skip_patterns = ["工 程 部 位", "强度等 级", "配合比编号", "试件编号"]
    data_rows = []
    for i in range(header_row_idx + 1, len(all_tables_data)):
        row = all_tables_data[i]
        first = row[0].strip()
        if not first or first in skip_patterns:
            continue
        if "工程名称" in first or "检测日期" in first:
            continue
        data_rows.append(row)

    if not data_rows:
        return None

    max_cols = max(len(r) for r in data_rows)
    aligned = []
    for row in data_rows:
        if len(row) < max_cols:
            row = row + [""] * (max_cols - len(row))
        aligned.append(row[:max_cols])

    df = pd.DataFrame(aligned, columns=header[:max_cols])

    col_map = {}
    for col in df.columns:
        if "浇筑日期" in col: col_map[col] = "浇筑日期"
        elif "制作日期" in col: col_map[col] = "浇筑日期"
        elif "龄期" in col: col_map[col] = "龄期"
        elif "等级" in col: col_map[col] = "等级"
        elif "工地" in col or "项目" in col: col_map[col] = "工地名称"
        elif "浇筑" in col or "部位" in col: col_map[col] = "浇筑部位"
    df.rename(columns=col_map, inplace=True)

    return df


def read_doc(path):
    """读取旧版 .doc 文件（需要 Windows + Word）"""
    try:
        import win32com.client, pythoncom, tempfile, os
        pythoncom.CoInitialize()
        try:
            wd = win32com.client.Dispatch("Word.Application")
            wd.Visible = False
            try:
                doc = wd.Documents.Open(os.path.abspath(path), ReadOnly=True)
                tmp = tempfile.NamedTemporaryFile(suffix=".txt", delete=False, encoding="gbk")
                tmp_path = tmp.name
                tmp.close()
                doc.SaveAs(tmp_path, FileFormat=2)
                doc.Close(False)
                with open(tmp_path, "r", encoding="gbk", errors="ignore") as f:
                    text = f.read()
                os.unlink(tmp_path)

                # 解析导出的文本
                lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
                header_idx = -1
                for i, line in enumerate(lines):
                    if any(x in line for x in ["工 程 部 位", "龄期"]):
                        header_idx = i
                        break

                if header_idx < 0:
                    return None

                data_rows = []
                for line in lines[header_idx:]:
                    cols = [c.strip() for c in line.split("|")]
                    if len(cols) < 3:
                        cols = [c.strip() for c in line.split("\t")]
                    if len(cols) >= 3:
                        first = cols[0].strip()
                        if first and first[0] in "1234567890":
                            data_rows.append(cols)

                if not data_rows:
                    return None

                max_cols = max(len(row) for row in data_rows)
                normalized = []
                for row in data_rows:
                    if len(row) < max_cols:
                        row = row + [""] * (max_cols - len(row))
                    normalized.append(row[:max_cols])

                headers = normalized[0]
                df = pd.DataFrame(normalized[1:], columns=headers)

                col_map = {}
                for col in df.columns:
                    if "浇筑日期" in col: col_map[col] = "浇筑日期"
                    elif "制作日期" in col: col_map[col] = "浇筑日期"
                    elif "龄期" in col: col_map[col] = "龄期"
                    elif "等级" in col: col_map[col] = "等级"
                    elif "工地" in col or "项目" in col: col_map[col] = "工地名称"
                    elif "浇筑" in col or "部位" in col: col_map[col] = "浇筑部位"
                df.rename(columns=col_map, inplace=True)

                return df
            finally:
                wd.Quit()
        finally:
            pythoncom.CoUninitialize()
    except Exception:
        pass

    # 方法2: LibreOffice 转换
    try:
        import subprocess, tempfile, os, shutil
        tmp_dir = tempfile.mkdtemp()
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx",
             "--outdir", tmp_dir, os.path.abspath(path)],
            capture_output=True, timeout=60
        )
        base = os.path.splitext(os.path.basename(path))[0]
        docx_path = os.path.join(tmp_dir, base + ".docx")
        if os.path.exists(docx_path):
            df = read_docx(docx_path)
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return df
    except Exception:
        pass

    return None


def filter_data(df, date_start, date_end, age_min, age_max, site_name, grade, pour_part):
    """过滤数据"""
    col_map = {col: col.strip().replace(" ", "").replace("\n", "") for col in df.columns}
    df = df.rename(columns=col_map)

    date_col = next((c for c in df.columns if "浇筑日期" in c), None)
    if not date_col:
        date_col = next((c for c in df.columns if "制作" in c and "日期" in c), None)

    age_col = next((c for c in df.columns if "龄期" in c), None)
    site_col = next((c for c in df.columns if "工地" in c), None)
    grade_col = next((c for c in df.columns if "等级" in c), None)
    pour_col = next((c for c in df.columns if "浇筑" in c), None)

    results = []
    for idx, row in df.iterrows():
        # 跳过表头行
        first = str(row.iloc[0]).strip() if len(row) > 0 else ""
        if first in ["序号", "序列", "NO", "编号", ""]:
            continue

        # 日期过滤
        if date_col:
            date_str = str(row[date_col]).strip()
            make_date = parse_date(date_str)
            if make_date:
                start_dt = parse_date(date_start) if date_start else None
                end_dt = parse_date(date_end) if date_end else None
                if start_dt and make_date < start_dt:
                    continue
                if end_dt and make_date > end_dt:
                    continue

        # 龄期过滤：优先从龄期列提取数字
        age_days = None
        if age_col:
            age_str = str(row[age_col]).strip()
            nums = re.findall(r'\d+', age_str)
            if nums:
                age_days = int(nums[0])

        # 没有龄期列则用日期计算
        if age_days is None and date_col:
            age_days = calc_age(str(row[date_col]).strip())

        if age_days is None:
            continue

        if not (age_min <= age_days <= age_max):
            continue

        if site_name and site_col:
            if site_name not in str(row[site_col]).strip():
                continue
        if grade and grade_col:
            if grade not in str(row[grade_col]).strip():
                continue
        if pour_part and pour_col:
            if pour_part not in str(row[pour_col]).strip():
                continue

        results.append((row, age_days))

    return results


def to_docx(results, output_path, all_columns, date_col, age_col, site_col, grade_col, pour_col):
    """将筛选结果写入 Word 文档"""
    doc = Document()
    title = doc.add_heading("混凝土试件记录筛选结果", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not results:
        doc.add_paragraph("（无符合条件的数据）")
        doc.save(output_path)
        return

    # 写入日期
    today_p = doc.add_paragraph(f"筛选日期：{datetime.now().strftime('%Y年%m月%d日')}")
    today_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 写入表头
    headers = ["序号", "浇筑日期", "龄期(天)", "等级", "工地名称", "浇筑部位"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for seq, (row, age) in enumerate(results, 1):
        cells = table.add_row().cells
        cells[0].text = str(seq)
        cells[1].text = str(row[date_col]).strip() if date_col and pd.notna(row[date_col]) else ""
        cells[2].text = str(age)
        cells[3].text = str(row[grade_col]).strip() if grade_col and pd.notna(row[grade_col]) else ""
        cells[4].text = str(row[site_col]).strip() if site_col and pd.notna(row[site_col]) else ""
        cells[5].text = str(row[pour_col]).strip() if pour_col and pd.notna(row[pour_col]) else ""

    doc.save(output_path)


class ConcreteFilterApp:
    def __init__(self, root):
        self.root = root
        root.title("混凝土试件记录筛选工具 v10")
        root.geometry("700x620")
        root.resizable(True, True)

        self.build_ui()

    def build_ui(self):
        f = ttk.Frame(self.root, padding=15)
        f.pack(fill="both", expand=True)

        # --- 文件选择区 ---
        file_frame = ttk.LabelFrame(f, text="文件/文件夹选择", padding=10)
        file_frame.pack(fill="x", pady=(0, 10))

        ttk.Label(file_frame, text="数据文件/文件夹：").grid(row=0, column=0, sticky="w")
        self.input_file = tk.StringVar()
        ttk.Entry(file_frame, textvariable=self.input_file, width=50).grid(row=0, column=1, sticky="ew", padx=(5, 5))
        ttk.Button(file_frame, text="浏览...", command=self.browse_input).grid(row=0, column=2)
        file_frame.columnconfigure(1, weight=1)

        ttk.Label(file_frame, text="输出文件：").grid(row=1, column=0, sticky="w", pady=(5, 0))
        self.output_file = tk.StringVar(value="筛选结果.docx")
        ttk.Entry(file_frame, textvariable=self.output_file, width=50).grid(row=1, column=1, sticky="ew", padx=(5, 5), pady=(5, 0))
        ttk.Button(file_frame, text="浏览...", command=self.browse_output).grid(row=1, column=2, pady=(5, 0))

        # --- 筛选条件区 ---
        cond_frame = ttk.LabelFrame(f, text="筛选条件", padding=10)
        cond_frame.pack(fill="x", pady=(0, 10))

        # 第一行：龄期
        ttk.Label(cond_frame, text="龄期范围（天）：").grid(row=0, column=0, sticky="w")
        self.age_min = tk.StringVar(value=str(DEFAULT_AGE_MIN))
        self.age_max = tk.StringVar(value=str(DEFAULT_AGE_MAX))
        ttk.Entry(cond_frame, textvariable=self.age_min, width=8).grid(row=0, column=1, sticky="w", padx=(5, 2))
        ttk.Label(cond_frame, text=" 至 ").grid(row=0, column=1, sticky="w", padx=(80, 2))
        ttk.Entry(cond_frame, textvariable=self.age_max, width=8).grid(row=0, column=1, sticky="e", padx=(90, 0))

        # 第二行：工地名称
        ttk.Label(cond_frame, text="工地名称：").grid(row=1, column=0, sticky="w", pady=(5, 0))
        self.site_name = tk.StringVar(value=DEFAULT_SITE_NAME)
        ttk.Entry(cond_frame, textvariable=self.site_name, width=30).grid(row=1, column=1, sticky="w", padx=(5, 0), pady=(5, 0))

        # 第三行：等级
        ttk.Label(cond_frame, text="强度等级：").grid(row=2, column=0, sticky="w", pady=(5, 0))
        self.grade = tk.StringVar(value=DEFAULT_GRADE)
        ttk.Entry(cond_frame, textvariable=self.grade, width=30).grid(row=2, column=1, sticky="w", padx=(5, 0), pady=(5, 0))

        # 第四行：浇筑部位
        ttk.Label(cond_frame, text="浇筑部位：").grid(row=3, column=0, sticky="w", pady=(5, 0))
        self.pour_part = tk.StringVar(value=DEFAULT_POUR_PART)
        ttk.Entry(cond_frame, textvariable=self.pour_part, width=30).grid(row=3, column=1, sticky="w", padx=(5, 0), pady=(5, 0))

        # 第五行：日期范围
        ttk.Label(cond_frame, text="制作日期范围：").grid(row=4, column=0, sticky="w", pady=(5, 0))
        self.date_start = tk.StringVar(value=DEFAULT_DATE_START)
        self.date_end = tk.StringVar(value=DEFAULT_DATE_END)
        ttk.Entry(cond_frame, textvariable=self.date_start, width=12).grid(row=4, column=1, sticky="w", padx=(5, 2), pady=(5, 0))
        ttk.Label(cond_frame, text=" 至 ").grid(row=4, column=1, sticky="w", padx=(120, 2), pady=(5, 0))
        ttk.Entry(cond_frame, textvariable=self.date_end, width=12).grid(row=4, column=1, sticky="e", padx=(145, 0), pady=(5, 0))
        ttk.Label(cond_frame, text="（格式：YYYY-MM-DD，留空表示不限）", foreground="gray").grid(row=5, column=0, columnspan=2, sticky="w", padx=(5, 0))

        # --- 日志区 ---
        log_frame = ttk.LabelFrame(f, text="处理日志", padding=10)
        log_frame.pack(fill="both", expand=True, pady=(0, 10))

        self.log_text = tk.Text(log_frame, height=12, state="disabled", wrap="word")
        scroll = ttk.Scrollbar(log_frame, command=self.log_text.yview)
        self.log_text.config(yscrollcommand=scroll.set)
        self.log_text.pack(side="left", fill="both", expand=True)
        scroll.pack(side="right", fill="y")

        # --- 运行按钮 ---
        btn_frame = ttk.Frame(f)
        btn_frame.pack(fill="x")
        self.btn_run = ttk.Button(btn_frame, text="▶ 开始筛选", command=self.run_filter)
        self.btn_run.pack(side="left")

    def log(self, msg):
        self.log_text.config(state="normal")
        self.log_text.insert("end", msg + "\n")
        self.log_text.see("end")
        self.log_text.config(state="disabled")
        self.root.update_idletasks()

    def browse_input(self):
        path = filedialog.askopenfilename(title="选择数据文件", filetypes=[
            ("支持格式", "*.xlsx *.xls *.csv *.docx *.doc"),
            ("所有文件", "*.*")
        ])
        if not path:
            path = filedialog.askdirectory(title="或选择文件夹")
        if path:
            self.input_file.set(path)
            if not self.output_file.get() or self.output_file.get() == "筛选结果.docx":
                base = os.path.splitext(os.path.basename(path))[0]
                self.output_file.set(f"{base}_筛选结果.docx")

    def browse_output(self):
        fname = filedialog.asksaveasfilename(
            title="保存结果文件",
            defaultextension=".docx",
            filetypes=[("Word文件", "*.docx")]
        )
        if fname:
            self.output_file.set(fname)

    def process_file(self, fpath, date_start, date_end, age_min, age_max, site_name, grade, pour_part):
        ext = os.path.splitext(fpath)[1].lower()
        if ext in [".xlsx", ".xls"]:
            df = read_excel(fpath)
        elif ext == ".csv":
            df = read_csv(fpath)
        elif ext == ".docx":
            df = read_docx(fpath)
        elif ext == ".doc":
            df = read_doc(fpath)
            if df is None:
                return None, None
        else:
            return None, None

        col_map = {col: col.strip().replace(" ", "").replace("\n", "") for col in df.columns}
        df = df.rename(columns=col_map)

        results = filter_data(df, date_start, date_end, age_min, age_max, site_name, grade, pour_part)
        return results, df.columns

    def run_filter(self):
        input_path = self.input_file.get().strip()
        output_path = self.output_file.get().strip()

        if not input_path:
            messagebox.showerror("错误", "请选择数据文件或文件夹！")
            return

        if not os.path.exists(input_path):
            messagebox.showerror("错误", f"路径不存在：{input_path}")
            return

        self.btn_run.config(state="disabled", text="处理中...")
        self.log_text.config(state="normal")
        self.log_text.delete("1.0", "end")
        self.log_text.config(state="disabled")

        try:
            date_start = self.date_start.get().strip()
            date_end = self.date_end.get().strip()
            try:
                age_min = int(self.age_min.get())
                age_max = int(self.age_max.get())
            except ValueError:
                messagebox.showerror("错误", "龄期必须是整数！")
                self.btn_run.config(state="normal", text="▶ 开始筛选")
                return
            site_name = self.site_name.get().strip()
            grade = self.grade.get().strip()
            pour_part = self.pour_part.get().strip()

            if os.path.isdir(input_path):
                self.log(f"📁 文件夹：{input_path}")
                supported_ext = [".xlsx", ".xls", ".csv", ".docx", ".doc"]
                files = []
                for root, dirs, files_in_dir in os.walk(input_path):
                    for f in files_in_dir:
                        if any(f.lower().endswith(ext) for ext in supported_ext):
                            files.append(os.path.join(root, f))

                if not files:
                    raise Exception(f"文件夹及子文件夹中未找到支持的数据文件")

                self.log(f"   找到 {len(files)} 个数据文件\n")
                all_results = []
                all_columns = None

                for i, fpath in enumerate(files, 1):
                    self.log(f"[{i}/{len(files)}] {os.path.basename(fpath)}")
                    results, cols = self.process_file(fpath, date_start, date_end, age_min, age_max, site_name, grade, pour_part)
                    if results is None:
                        self.log(f"   ⚠️ 读取失败，跳过")
                        continue
                    self.log(f"   → {len(results)} 条记录")
                    all_results.extend(results)
                    if cols is not None:
                        all_columns = cols

                if all_results:
                    date_col = next((c for c in all_columns if "浇筑日期" in c), None) or next((c for c in all_columns if "制作" in c and "日期" in c), None)
                    age_col = next((c for c in all_columns if "龄期" in c), None)
                    site_col = next((c for c in all_columns if "工地" in c), None)
                    grade_col = next((c for c in all_columns if "等级" in c), None)
                    pour_col = next((c for c in all_columns if "浇筑" in c), None)
                    to_docx(all_results, output_path, all_columns, date_col, age_col, site_col, grade_col, pour_col)
                    self.log(f"\n✅ 完成！共 {len(all_results)} 条记录")
                    self.log(f"📄 结果已保存：{output_path}")
                    messagebox.showinfo("完成", f"筛选完成！\n\n共处理 {len(files)} 个文件\n合计找到 {len(all_results)} 条记录\n\n已保存到：{output_path}")
                else:
                    to_docx([], output_path, all_columns or [], None, None, None, None, None)
                    self.log(f"\n⚠️ 无符合条件的数据")
                    self.log(f"📄 结果已保存：{output_path}")
                    messagebox.showinfo("完成", f"筛选完成，但无符合条件的数据\n\n已保存到：{output_path}")

            else:
                self.log(f"📖 文件：{input_path}")
                results, cols = self.process_file(input_path, date_start, date_end, age_min, age_max, site_name, grade, pour_part)
                if results is None:
                    raise Exception("无法读取文件，请确认文件格式正确")
                self.log(f"   共 {len(results)} 条记录\n")

                if results:
                    date_col = next((c for c in cols if "浇筑日期" in c), None) or next((c for c in cols if "制作" in c and "日期" in c), None)
                    age_col = next((c for c in cols if "龄期" in c), None)
                    site_col = next((c for c in cols if "工地" in c), None)
                    grade_col = next((c for c in cols if "等级" in c), None)
                    pour_col = next((c for c in cols if "浇筑" in c), None)
                    to_docx(results, output_path, cols, date_col, age_col, site_col, grade_col, pour_col)
                    self.log(f"✅ 完成！")
                    self.log(f"📄 结果已保存：{output_path}")
                    messagebox.showinfo("完成", f"筛选完成！\n\n共找到 {len(results)} 条记录\n\n已保存到：{output_path}")
                else:
                    to_docx([], output_path, cols or [], None, None, None, None, None)
                    self.log(f"⚠️ 无符合条件的数据")
                    self.log(f"📄 结果已保存：{output_path}")
                    messagebox.showinfo("完成", f"筛选完成，但无符合条件的数据\n\n已保存到：{output_path}")

        except Exception as e:
            messagebox.showerror("错误", str(e))
            self.log(f"\n❌ 错误：{e}")
        finally:
            self.btn_run.config(state="normal", text="▶ 开始筛选")


def main():
    root = tk.Tk()
    app = ConcreteFilterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
